"""
jd_parser.py
============
Extracts plain text from an uploaded job description.

Companies send JDs as PDF or Word attachments far more often than as pasteable
text, so the tailoring flow takes a file rather than a textarea. Supported:
``.pdf``, ``.docx``, ``.txt`` and ``.md``.

Deliberately not supported
--------------------------
* **Legacy ``.doc``** (Word 97-2003) is a compound binary format that needs
  LibreOffice or antiword to read. Rejected with an instruction to re-save,
  rather than dragging in a system dependency.
* **Scanned PDFs** produce no extractable text. We detect the empty result and
  say so explicitly - OCR is a much heavier dependency than this feature earns.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional

logger = logging.getLogger("resumemaxxer.jd_parser")

# Guard rails. A job description is a few kilobytes of text; anything far
# larger is a mistake or an attack, and parsing it wastes memory and time.
MAX_UPLOAD_BYTES = 5 * 1024 * 1024  # 5 MB
MAX_EXTRACTED_CHARS = 20_000
MIN_EXTRACTED_CHARS = 50

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


class JDParseError(ValueError):
    """The upload could not be turned into usable text.

    Routers map this to HTTP 400 - it is always the user's file that is the
    problem, never the server.
    """


def _clean(text: str) -> str:
    """Normalise extracted text into something worth sending to Gemini.

    PDF extraction in particular produces ragged output: hard-wrapped lines,
    runs of blank lines from layout boxes, and non-breaking spaces. Tidying it
    up cuts token count and materially improves keyword extraction.
    """
    # Normalise unicode spaces and bullet glyphs that survive extraction.
    text = text.replace(" ", " ").replace("•", "- ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse runs of spaces/tabs, but keep newlines - the line structure of a
    # JD carries meaning (requirement lists, section headers).
    text = re.sub(r"[ \t]+", " ", text)
    # Three or more newlines become two: keeps paragraphs, drops page gaps.
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.strip() for line in text.split("\n"))

    return text.strip()


def _extract_pdf(data: bytes) -> str:
    """Pull text out of a PDF, page by page."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise JDParseError(
            "PDF support is not installed on the server (pypdf missing)."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise JDParseError(
            "That PDF could not be opened. It may be corrupted or password-protected."
        ) from exc

    if reader.is_encrypted:
        # Some PDFs are encrypted with an empty owner password, which pypdf can
        # open; try that before giving up.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise JDParseError(
                "That PDF is password-protected. Remove the password and retry."
            ) from exc

    pages: list[str] = []
    for index, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # One unreadable page should not lose the whole document.
            logger.warning("Could not extract page %d of the uploaded PDF", index + 1)

    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    """Pull text out of a .docx, including tables.

    Tables matter here: JDs routinely lay out requirements in a two-column
    table, and paragraph-only extraction would silently drop half the content.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise JDParseError(
            "DOCX support is not installed on the server (python-docx missing)."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise JDParseError(
            "That .docx could not be opened. If it is an older .doc file, "
            "open it in Word and re-save as .docx."
        ) from exc

    parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Word repeats the cell object across merged spans; de-duplicate
            # consecutive identical cells so merged headers appear once.
            deduped: list[str] = []
            for cell in cells:
                if cell and (not deduped or deduped[-1] != cell):
                    deduped.append(cell)
            if deduped:
                parts.append(" | ".join(deduped))

    return "\n".join(parts)


def _extract_plaintext(data: bytes) -> str:
    """Decode a .txt or .md upload, tolerating unknown encodings."""
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 cannot actually fail, but be explicit rather than clever.
    return data.decode("utf-8", errors="replace")


def extract_text(filename: Optional[str], data: bytes) -> str:
    """Turn an uploaded job description into clean plain text.

    Raises `JDParseError` with a message written for the student, since it is
    surfaced directly in the UI.
    """
    if not data:
        raise JDParseError("That file is empty.")

    if len(data) > MAX_UPLOAD_BYTES:
        size_mb = len(data) / (1024 * 1024)
        raise JDParseError(
            f"That file is {size_mb:.1f} MB. The limit is "
            f"{MAX_UPLOAD_BYTES // (1024 * 1024)} MB."
        )

    name = (filename or "").lower().strip()

    if name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".docx"):
        text = _extract_docx(data)
    elif name.endswith((".txt", ".md")):
        text = _extract_plaintext(data)
    elif name.endswith(".doc"):
        raise JDParseError(
            "Legacy .doc files are not supported. Open it in Word or Google "
            "Docs and save it as .docx or PDF."
        )
    else:
        raise JDParseError(
            "Unsupported file type. Upload a PDF, DOCX, TXT or MD file."
        )

    text = _clean(text)

    if len(text) < MIN_EXTRACTED_CHARS:
        # Almost always a scanned/image-only PDF.
        raise JDParseError(
            "Almost no text could be read from that file. If it is a scanned "
            "PDF or an image, the text is not selectable - copy the job "
            "description into a .txt file and upload that instead."
        )

    if len(text) > MAX_EXTRACTED_CHARS:
        # Keep the head: JDs put the role and requirements up top, and boilerplate
        # (equal-opportunity statements, benefits) at the bottom.
        logger.info("Truncating extracted JD from %d chars", len(text))
        text = text[:MAX_EXTRACTED_CHARS]

    return text
