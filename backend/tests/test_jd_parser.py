"""
Job-description parsing.

Every branch here is a file a student will actually upload, and the failure
modes matter: a silently bad parse produces a plausible but generic resume,
which is worse than an error.
"""

from __future__ import annotations

import io

import pytest

import jd_parser

JD_TEXT = """Software Engineering Intern - Backend
Bengaluru, India

Requirements:
- Strong Python fundamentals; familiarity with FastAPI
- Working knowledge of PostgreSQL and writing efficient SQL
- Comfortable with Git, Docker and CI/CD pipelines
"""


@pytest.fixture(scope="module")
def pdf_bytes() -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    y = 800
    for line in JD_TEXT.split("\n"):
        pdf.drawString(50, y, line)
        y -= 14
    pdf.save()
    return buffer.getvalue()


@pytest.fixture(scope="module")
def docx_bytes() -> bytes:
    """A .docx whose requirements live in a table.

    This is the case that matters: JDs routinely use a table, and
    paragraph-only extraction would silently drop half the document.
    """
    import docx

    document = docx.Document()
    document.add_paragraph("Software Engineering Intern - Backend")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Must have"
    table.rows[0].cells[1].text = "Python, FastAPI, PostgreSQL"
    table.rows[1].cells[0].text = "Tooling"
    table.rows[1].cells[1].text = "Git, Docker, CI/CD"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extracts_text_from_pdf(pdf_bytes):
    text = jd_parser.extract_text("jd.pdf", pdf_bytes)
    assert "FastAPI" in text
    assert "PostgreSQL" in text


def test_extracts_docx_including_table_rows(docx_bytes):
    text = jd_parser.extract_text("jd.docx", docx_bytes)
    assert "Tooling | Git, Docker, CI/CD" in text


@pytest.mark.parametrize(
    "filename, encoding",
    [("jd.txt", "utf-8"), ("jd.md", "utf-16")],
)
def test_extracts_plaintext_in_several_encodings(filename, encoding):
    text = jd_parser.extract_text(filename, JD_TEXT.encode(encoding))
    assert "FastAPI" in text


@pytest.mark.parametrize(
    "filename, data, expected_message",
    [
        pytest.param("jd.doc", b"\xd0\xcf\x11\xe0" + b"x" * 200,
                     "Legacy .doc", id="legacy-doc"),
        pytest.param("jd.exe", b"MZ" + b"x" * 200,
                     "Unsupported file type", id="unsupported-type"),
        pytest.param("jd.pdf", b"", "empty", id="empty-file"),
        pytest.param("jd.pdf", b"x" * (jd_parser.MAX_UPLOAD_BYTES + 1),
                     "limit is", id="too-large"),
        pytest.param("jd.txt", b"hi", "Almost no text", id="scanned-or-near-empty"),
        pytest.param("jd.pdf", b"%PDF-1.4 garbage" + b"\x00" * 300,
                     "could not be opened", id="corrupt-pdf"),
    ],
)
def test_bad_uploads_are_rejected_with_a_useful_message(filename, data, expected_message):
    with pytest.raises(jd_parser.JDParseError) as exc:
        jd_parser.extract_text(filename, data)
    assert expected_message in str(exc.value)


def test_long_documents_are_truncated_not_rejected():
    """Boilerplate lives at the bottom of a JD, so keeping the head is right."""
    padding = ("Python FastAPI PostgreSQL requirement line.\n"
               * (jd_parser.MAX_EXTRACTED_CHARS // 20))
    text = jd_parser.extract_text("jd.txt", padding.encode())
    assert len(text) == jd_parser.MAX_EXTRACTED_CHARS
